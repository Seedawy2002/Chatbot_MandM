import streamlit as st
import os
import tempfile
import pickle
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import shutil
from pdfminer.high_level import extract_text
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import time
from docx import Document as DocxDocument

# Set page config at the very beginning
st.set_page_config(page_title="CV Management System", layout="wide")

# Add custom CSS for styling messages
st.markdown(
    """
    <style>
    .user-message {
        background-color: #1976d2;
        border-left: 4px solid #e1f5fe;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .assistant-message {
        background-color: #6c19d2;
        border-left: 4px solid #f3e1fe;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    </style>
    """,
    unsafe_allow_html=True
)

class CVManager:
    def __init__(self, resumes_dir: str = 'resumes'):
        self.resumes_dir = resumes_dir
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        
        # Create resumes directory if it doesn't exist
        os.makedirs(self.resumes_dir, exist_ok=True)
        
        # Check if preprocessing is needed
        self.check_and_run_preprocessing()

    def check_and_run_preprocessing(self):
        """
        Check if the required files (chunks.pkl, embeddings.pkl, resumes_index.faiss) exist.
        If not, run the preprocessing steps.
        """
        required_files = ['chunks.pkl', 'embeddings.pkl', 'resumes_index.faiss']
        if not all(os.path.exists(file) for file in required_files):
            st.info("Preprocessing resumes... This may take a few minutes.")
            self.process_resumes()
            self.generate_embeddings()
            self.create_faiss_index()
            st.success("Preprocessing completed!")

    def process_resumes(self):
        """
        Process resumes in the resumes directory and save chunks to chunks.pkl.
        """
        pdf_files = [f for f in os.listdir(self.resumes_dir) if f.endswith('.pdf') or f.endswith('.docx')]
        chunks = []

        for pdf_file in pdf_files:
            pdf_path = os.path.join(self.resumes_dir, pdf_file)
            try:
                text = self.extract_text_from_file(pdf_path)
                doc = Document(page_content=text, metadata={'source': pdf_file})
                chunked_docs = self.text_splitter.split_documents([doc])
                chunks.extend(chunked_docs)
            except Exception as e:
                st.error(f"Error processing {pdf_path}: {e}")
                continue

        with open('chunks.pkl', 'wb') as f:
            pickle.dump(chunks, f)

    def generate_embeddings(self):
        """
        Generate embeddings from chunks and save to embeddings.pkl.
        """
        with open('chunks.pkl', 'rb') as f:
            chunks = pickle.load(f)

        embeddings = [self.embedding_model.encode(chunk.page_content) for chunk in chunks]

        with open('embeddings.pkl', 'wb') as f:
            pickle.dump(embeddings, f)

    def create_faiss_index(self):
        """
        Create a FAISS index from embeddings and save to resumes_index.faiss.
        """
        with open('embeddings.pkl', 'rb') as f:
            embeddings = pickle.load(f)

        embeddings_array = np.array(embeddings).astype('float32')
        dimension = embeddings_array.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings_array)

        faiss.write_index(index, 'resumes_index.faiss')

    def extract_text_from_file(self, file_path):
        if file_path.endswith('.pdf'):
            return extract_text(file_path)
        elif file_path.endswith('.docx'):
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Unsupported file format")

    def add_new_cvs(self, cv_files) -> bool:
        try:
            for cv_file in cv_files:
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(cv_file.name)[1]) as tmp_file:
                    tmp_file.write(cv_file.getvalue())
                    temp_path = tmp_file.name

                cv_filename = cv_file.name
                destination = os.path.join(self.resumes_dir, cv_filename)
                shutil.copy2(temp_path, destination)
                os.unlink(temp_path)
            
            # Re-run preprocessing to update the index
            self.check_and_run_preprocessing()
            return True
            
        except Exception as e:
            st.error(f"Error adding CVs: {str(e)}")
            return False

    def delete_cv(self, cv_filename):
        try:
            os.remove(os.path.join(self.resumes_dir, cv_filename))
            # Re-run preprocessing to update the index
            self.check_and_run_preprocessing()
            return True
        except Exception as e:
            st.error(f"Error deleting CV: {str(e)}")
            return False

    def search_resumes(self, query: str, k: int = 5):
        try:
            index = faiss.read_index('resumes_index.faiss')
            with open('chunks.pkl', 'rb') as f:
                chunks = pickle.load(f)

            query_embedding = self.embedding_model.encode(query)
            query_embedding = np.array([query_embedding]).astype('float32')

            distances, indices = index.search(query_embedding, k)

            max_distance = np.max(distances)
            confidence_scores = 1 - (distances / max_distance)

            results = []
            for i, idx in enumerate(indices[0]):
                result = {
                    "text": chunks[idx].page_content,
                    "source": chunks[idx].metadata['source'],
                    "confidence": float(confidence_scores[0][i])
                }
                results.append(result)

            return results
        except Exception as e:
            st.error(f"Error searching resumes: {str(e)}")
            return []

def init_session_state():
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'upload_time' not in st.session_state:
        st.session_state.upload_time = None
    if 'status_message' not in st.session_state:
        st.session_state.status_message = None
    if 'status_time' not in st.session_state:
        st.session_state.status_time = None
    if 'current_upload' not in st.session_state:
        st.session_state.current_upload = None
    if 'selected_cv' not in st.session_state:
        st.session_state.selected_cv = None
    if 'search_results' not in st.session_state:
        st.session_state.search_results = None

def check_time_elapsed():
    current_time = time.time()
    
    if st.session_state.status_time is not None:
        if current_time - st.session_state.status_time > 30:
            st.session_state.status_message = None
            st.session_state.status_time = None
    
    if st.session_state.upload_time is not None:
        if current_time - st.session_state.upload_time > 30:
            st.session_state.current_upload = None
            st.session_state.upload_time = None

def get_chatbot_response(query, context, cv_manager):
    results = cv_manager.search_resumes(query)
    
    if not results:
        return "I couldn't find any relevant information about that. Could you please rephrase your question?"

    response = ""
    confidence_sum = 0
    
    for result in results:
        if result["confidence"] > 0.5:
            response += f"{result['text']} "
            confidence_sum += result["confidence"]
    
    if not response:
        return "I found some information but I'm not confident enough about its relevance. Could you please be more specific?"

    avg_confidence = confidence_sum / len(results)
    
    if "experience" in query.lower():
        response = f"Based on the CV information (confidence: {avg_confidence:.2f}), the candidate has the following experience: {response}"
    elif "skills" in query.lower():
        response = f"According to the CV (confidence: {avg_confidence:.2f}), the candidate's skills include: {response}"
    elif "education" in query.lower():
        response = f"The educational background (confidence: {avg_confidence:.2f}) shows: {response}"
    else:
        response = f"Here's what I found (confidence: {avg_confidence:.2f}): {response}"

    return response

def main():
    st.title("CV Management System")
    
    init_session_state()
    
    cv_manager = CVManager()
    
    page = st.sidebar.selectbox("Choose a page", ["Upload CV", "CV Chatbot", "View All CVs", "Search CV"])
    
    if page == "Upload CV":
        st.header("Upload New CV")
        
        upload_placeholder = st.empty()
        status_placeholder = st.empty()
        
        check_time_elapsed()
        
        if st.session_state.status_message:
            status_placeholder.success(st.session_state.status_message)
            time.sleep(10)
            st.session_state.status_message = None
            st.rerun()
        
        uploaded_files = upload_placeholder.file_uploader("Choose PDF or DOCX files", type=['pdf', 'docx'], accept_multiple_files=True)
        
        if uploaded_files:
            if st.button("Process CVs"):
                with st.spinner("Processing CVs..."):
                    success = cv_manager.add_new_cvs(uploaded_files)
                    if success:
                        st.session_state.status_message = "✅ Successfully added CVs"
                        st.session_state.status_time = time.time()
                    else:
                        st.session_state.status_message = "❌ Failed to process CVs"
                        st.session_state.status_time = time.time()
                    
                    st.session_state.current_upload = None
                    st.rerun()
    
    elif page == "CV Chatbot":
        st.header("CV Chatbot")
        
        chat_container = st.container()
        with chat_container:
            for message in st.session_state.chat_history:
                if message["role"] == "user":
                    st.markdown(f'<div class="user-message">You: {message["content"]}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="assistant-message">Assistant: {message["content"]}</div>', unsafe_allow_html=True)
        
        user_input = st.text_input("Ask about candidates:", key="user_input")
        
        if st.button("Send") and user_input:
            st.session_state.chat_history.append({"role": "user", "content": user_input})
            response = get_chatbot_response(user_input, st.session_state.chat_history, cv_manager)
            st.session_state.chat_history.append({"role": "assistant", "content": response})
            st.rerun()
    
    elif page == "View All CVs":
        st.header("All CVs")
        
        resumes = [f for f in os.listdir(cv_manager.resumes_dir) if f.endswith('.pdf') or f.endswith('.docx')]
        if resumes:
            for resume in resumes:
                col1, col2 = st.columns([0.8, 0.2])
                with col1:
                    st.text(resume)
                with col2:
                    if st.button(f"Delete {resume}"):
                        if cv_manager.delete_cv(resume):
                            st.success(f"Deleted {resume}")
                            st.rerun()
                        else:
                            st.error(f"Failed to delete {resume}")
        else:
            st.info("No CVs found in the system")
    
    elif page == "Search CV":
        st.header("Search CV")
        
        # User input for search query
        search_query = st.text_input("Enter your search query (e.g., skills, experience):")
        
        if search_query:
            num_results = st.slider("Select the number of results to display", 1, 10, 5)
            
            if st.button("Search"):
                with st.spinner("Searching CVs..."):
                    search_results = cv_manager.search_resumes(search_query, k=num_results)
                    
                    if search_results:
                        st.subheader("Search Results")
                        for i, result in enumerate(search_results):
                            st.write(f"**Result {i+1}:** {result['source']} (Confidence: {result['confidence']:.2f})")
                            st.write(f"**Relevant Text:** {result['text']}")
                            st.write("---")
                        
                        # Store search results in session state
                        st.session_state.search_results = search_results
                        
                        # Get the list of CV sources
                        cv_sources = [result['source'] for result in search_results]
                        
                        # Initialize selected_cv if not already set or if it's not in the current results
                        if 'selected_cv' not in st.session_state or st.session_state.selected_cv not in cv_sources:
                            st.session_state.selected_cv = cv_sources[0]  # Default to the first CV
                        
                        # Allow user to select a CV
                        selected_cv = st.selectbox(
                            "Select a CV to view details:",
                            cv_sources,
                            index=cv_sources.index(st.session_state.selected_cv)
                        )
                        
                        # Update selected CV in session state
                        st.session_state.selected_cv = selected_cv
                        
                        # Display selected CV details
                        if st.session_state.selected_cv:
                            st.subheader(f"Selected CV: {st.session_state.selected_cv}")
                            cv_path = os.path.join(cv_manager.resumes_dir, st.session_state.selected_cv)
                            cv_text = cv_manager.extract_text_from_file(cv_path)
                            st.write(cv_text)
                            
                            # Chatbot for selected CV
                            st.subheader("Ask about this candidate:")
                            user_input = st.text_input("Ask a question (e.g., skills, experience):")
                            
                            if st.button("Ask"):
                                if user_input:
                                    response = get_chatbot_response(user_input, st.session_state.chat_history, cv_manager)
                                    st.write(f"**Assistant:** {response}")
                    else:
                        st.info("No relevant CVs found for the given query.")

if __name__ == "__main__":
    main()