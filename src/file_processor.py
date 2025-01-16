import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import FastEmbedEmbeddings
from langchain.vectorstores.utils import filter_complex_metadata
from langchain_community.document_loaders import UnstructuredFileLoader
from langchain_pinecone import Pinecone
import uuid
from langchain.schema import Document
from docx import Document as DocxDocument
import spacy
import pdfplumber

os.environ.get("GROQ_API_KEY", "PINECONE_API_KEY")

class FileProcessor:
    def __init__(self, folder_path, index_name="new", embedding_model="BAAI/bge-base-en-v1.5", chunk_size=1000, chunk_overlap=100):
        self.folder_path = folder_path
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embeddings = FastEmbedEmbeddings(model_name=embedding_model)
        self.vector_db = Pinecone.from_existing_index(
            index_name=index_name,
            embedding=self.embeddings,
            text_key="text"
        )
        self.nlp = spacy.load("en_core_web_sm")

    def process_files(self):
        """Processes all supported files in the folder."""
        for file_name in os.listdir(self.folder_path):
            file_path = os.path.join(self.folder_path, file_name)
            if file_name.endswith('.pdf') or file_name.endswith('.docx'):
                self.prepare_file(file_path)

    def prepare_file(self, file_path):
        """Reads, splits, and stores embeddings for a single file."""
        pages = self.read_file(file_path)
        texts = self.split_document(pages)
        self.store_embeddings(file_path, texts)

    def read_file(self, file_path):
        """Reads the content of a file based on its extension."""
        if file_path.endswith('.pdf'):
            return self.read_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self.read_docx(file_path)
        else:
            raise ValueError("Unsupported file type.")

    def read_pdf(self, file_path):
        """Reads a PDF file and extracts its text."""
        with pdfplumber.open(file_path) as pdf:
            pages = [page.extract_text() for page in pdf.pages if page.extract_text()]
        return [Document(page_content=page) for page in pages if page]


    def read_docx(self, file_path):
        """Reads a DOCX file and returns Document objects."""
        doc = DocxDocument(file_path)
        return [Document(page_content=p.text) for p in doc.paragraphs if p.text.strip()]

    def split_document(self, pages):
        """Splits the document into smaller chunks."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap
        )
        return text_splitter.split_documents(pages)

    def extract_name(self, text):
        """Extracts the name of the person from the CV."""
        doc = self.nlp(text)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                return ent.text
        return "Unknown Name"

    def store_embeddings(self, file_path, texts):
        """Generates embeddings and stores them with file metadata."""
        # Combine all chunk content into a single text for name extraction
        full_text = " ".join([doc.page_content for doc in texts])
        extracted_name = self.extract_name(full_text)

        # Create metadata and text with embedded name
        metadatas = [
            {
                "source": file_path,
                "id": str(uuid.uuid4()),
                "person_name": extracted_name
            }
            for _ in texts
        ]
        text_contents = [
            f"### CV Analysis for {extracted_name} ###\n\n{doc.page_content}" for doc in texts
        ]
        # Add texts and metadata to vector store
        self.vector_db.add_texts(texts=text_contents, metadatas=metadatas)